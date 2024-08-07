import os
import docx
from langchain_groq import ChatGroq
from langchain_openai import ChatOpenAI
import streamlit as st  # Import Streamlit for web application interface
from docx import Document  # Import python-docx for Word document creation
from io import BytesIO  # Import BytesIO for in-memory file operations
import replicate  # Import Replicate for image generation
import requests  # Import requests to download images
import asyncio
import google.generativeai as genai  # Import the appropriate module for Gemini
from langchain_google_genai import ChatGoogleGenerativeAI
from crewai import Agent, Task, Crew
from langchain_community.tools import DuckDuckGoSearchRun


def verify_gemini_api_key(api_key):
    API_VERSION = 'v1'
    api_url = f"https://generativelanguage.googleapis.com/{API_VERSION}/models?key={api_key}"
    
    try:
        response = requests.get(api_url, headers={'Content-Type': 'application/json'})
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        # If we get here, it means the request was successful
        return True
    
    except requests.exceptions.HTTPError as e:
        
        return False
    
    except requests.exceptions.RequestException as e:
        # For any other request-related exceptions
        raise ValueError(f"An error occurred: {str(e)}")

def verify_gpt_api_key(api_key):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # Using a simple request to the models endpoint
    response = requests.get("https://api.openai.com/v1/models", headers=headers)
    
    if response.status_code == 200:
        return True
    elif response.status_code == 401:
        return False
    else:
        print(f"Unexpected status code: {response.status_code}")
        return False
    
def verify_groq_api_key(api_key):
    api_url = "https://api.groq.com/openai/v1/models"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    try:
        response = requests.get(api_url, headers=headers)
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        # If we get here, it means the request was successful
        return True
    
    except requests.exceptions.HTTPError as e:
        
        return False
    
    except requests.exceptions.RequestException as e:
        # For any other request-related exceptions
        raise ValueError(f"An error occurred: {str(e)}")
    
def verify_replicate_api_key(api_key):
    api_url = "https://api.replicate.com/v1/models"
    headers = {
        "Authorization": f"Token {api_key}",
        "Content-Type": "application/json"
    }
    
    try:
        response = requests.get(api_url, headers=headers)
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        # If we get here, it means the request was successful
        return True
    
    except requests.exceptions.HTTPError as e:
    
        return False
        
    except requests.exceptions.RequestException as e:
        # For any other request-related exceptions
        raise ValueError(f"An error occurred: {str(e)}")
    
def generate_text(llm, topic):
    
    search_tool = DuckDuckGoSearchRun(
    name="duckduckgo_search",
    description="""Search the web using DuckDuckGo. Give argument -
                {"query": "<Whatever you want to search>"}""",
)
    inputs = {'topic': topic}

    # Define Blog Researcher Agent
    blog_researcher = Agent(
        role='Blog Content Researcher',
        goal='Conduct thorough research to uncover compelling insights for engaging blog content.',
        backstory=("An experienced content strategist with a knack for analyzing trends and audience behavior, "
                   "delivering actionable insights for high-quality blog content."),
        verbose=True,
        allow_delegation=False,
        llm=llm
        )

    # Define Blog Writer Agent
    blog_writer = Agent(
        role='Blog Writer',
        goal='Craft authoritative and engaging blog content that resonates with the audience and establishes the brand as a leader.',
        backstory=("A seasoned writer known for distilling complex topics into captivating stories, with a deep understanding of audience psychology."),
        verbose=True,
        allow_delegation=False,
        llm=llm,
        max_iter=5

    )

    # Define Blog Reviewer Agent
    blog_reviewer = Agent(
        role='Content Reviewer',
        goal='Review and refine blog drafts to ensure they meet high standards of quality and impact.',
        backstory=("An expert editor with a meticulous eye for detail, known for elevating content to publication-ready standards."),
        verbose=True,
        allow_delegation=False,
        llm=llm,
        max_iter=5

    )

    # Define Task for Researcher
    task_researcher = Task(
        description=(f"Research the latest trends and insights on {topic}. Identify key developments, emerging trends, unique perspectives, and content ideas. You can use the serach tool if needed "),
        agent=blog_researcher,
        expected_output=(
            f"1. Overview and background of {topic}.\n"
            "2. Recent key developments.\n"
            "3. Emerging trends and innovative approaches.\n"
            "4. Unique angles and untapped opportunities.\n"
            "5. Potential content ideas with brief descriptions.\n"
            "6. List of relevant sources."
        ),
        tools=[search_tool]
        )

    # Define Task for Writer
    task_writer = Task(
        description=(f"Based on the research report, craft an engaging and authoritative blog post on {topic}."),
        agent=blog_writer,
        expected_output=(
            "1. Engaging introduction with a hook.\n"
            "2. Detailed exploration of key developments.\n"
            "3. Emerging trends and innovative ideas in content.\n"
            "4. Unique angles and perspectives in content.\n"
            "5. Clear explanations of complex concepts.\n"
            "6. Compelling conclusion.\n"
        )
    )

    # Define Task for Reviewer
    task_reviewer = Task(
        description=(f"Review the drafted blog content on {topic}, providing detailed feedback and revisions for quality and impact."),
        agent=blog_reviewer,
        expected_output=(
            "1. Overall content assessment.\n"
            "2. Identification of inaccuracies and gaps.\n"
            "3. Suggestions for improving flow and readability.\n"
            "4. Recommendations for tone and voice.\n"
            "5. Edits for grammar and punctuation.\n"
            "6. Final assessment of readiness."
        )
    )

    # Define Task for Final Writer
    task_final_writer = Task(
        description=(f"Revise the blog content on {topic} based on the reviewer's feedback, ensuring it meets high standards."),
        agent=blog_writer,
        expected_output=(
            "1. Factually accurate and corrected content.\n"
            "2. Clear, well-structured flow.\n"
            "3. Concise and engaging language.\n"
            "4. Consistent tone and voice.\n"
            "5. Enhanced insights.\n"
            "6. Addressed reviewer feedback.\n"
            "7. Creative and engaging blog title.\n"
            "8. Final draft of at least 1000 words.\n"
            "9. Don't include any agentic thoughts and just give a ready blog without any extra comments."
        )
    )

    # Initialize Crew: Coordinates agents and tasks for structured blog content workflow
    crew = Crew(
        agents=[blog_researcher, blog_writer, blog_reviewer, blog_writer],
        tasks=[task_researcher, task_writer, task_reviewer, task_final_writer],
        verbose=2,
        context={"Blog Topic is ": topic}
    )

    # Start the workflow and generate the result
    result = crew.kickoff(inputs=inputs)

    return result


# Function to generate images based on prompts
def generate_images(replicate_api_token, prompt):
    os.environ["REPLICATE_API_TOKEN"] = replicate_api_token

    # Define the input for the image generation
    input = {
        "prompt": prompt,
        "scheduler": "K_EULER"
    }

    # Generate the image
    output = replicate.run(
        "stability-ai/stable-diffusion:ac732df83cea7fff18b8472768c88ad041fa750ff7682a21affe81863cbe77e4",
        input=input
    )

    # Assuming output is a list of URLs, return the first one
    if output and isinstance(output, list) and len(output) > 0:
        return output[0]
    else:
        raise ValueError("No image URL returned from Replicate API.")

    
# Streamlit web application
def main():
    st.header('AI Blog Content Generator')
    validity_model= False
    validity_replicate = False

    # Initialize session state
    if 'generated_content' not in st.session_state:
        st.session_state.generated_content = None
    if 'generated_image_url' not in st.session_state:
        st.session_state.generated_image_url = None
    if 'topic' not in st.session_state:
        st.session_state.topic = ""

    with st.sidebar:
        with st.form('Gemini/OpenAI/Groq'):
            # User selects the model (Gemini/Cohere) and enters API keys
            model = st.radio('Choose Your LLM', ('Gemini', 'OpenAI', 'Groq'))
            api_key = st.text_input(f'Enter your API key', type="password")
            replicate_api_token = st.text_input('Enter Replicate API key', type="password")
            submitted = st.form_submit_button("Submit")
            
        if api_key and replicate_api_token:
            if model == "Gemini":
                validity_model = verify_gemini_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")
            elif model == "OpenAI":
                validity_model = verify_gpt_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")            
            elif model == "Groq":
                validity_model = verify_groq_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")
            
            validity_replicate = verify_replicate_api_key(replicate_api_token)
            if validity_replicate ==True:
                st.write(f"Valid Replicate API key")
            else:
                st.write(f"Invalid Replicate API key")
    
    if validity_model and validity_replicate:
        if model == 'OpenAI':
            async def setup_OpenAI():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                os.environ["OPENAI_API_KEY"] = api_key
                llm = ChatOpenAI(model='gpt-4-turbo', temperature=0.6, max_tokens=2000, api_key=api_key)
                return llm

            llm = asyncio.run(setup_OpenAI())

        elif model == 'Gemini':
            async def setup_gemini():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    verbose=True,
                    temperature=0.6,
                    google_api_key=api_key  # Use the API key from the environment variable
                )
                return llm

            llm = asyncio.run(setup_gemini())

        elif model == 'Groq':
            async def setup_groq():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGroq(
                    api_key=api_key,
                    model='llama3-70b-8192'
                )
                return llm

            llm = asyncio.run(setup_groq())


        topic = st.text_input("Enter the blog topic:")
        st.session_state.topic = topic

        if st.button("Generate Blog Content"):
            with st.spinner("Generating content..."):
                st.session_state.generated_content = generate_text(llm, st.session_state.topic)
                st.session_state.generated_image_url = generate_images(replicate_api_token, st.session_state.topic)

        # Display content if it exists in session state
        if st.session_state.generated_content and st.session_state.generated_image_url:
            content_lines = st.session_state.generated_content.split('\n')
            first_line = content_lines[0]
            remaining_content = '\n'.join(content_lines[1:])

            st.markdown(first_line)
            st.image(st.session_state.generated_image_url, caption="Generated Image", use_column_width=True)
            st.markdown(remaining_content)

            # Download the images and add them to the document
            response = requests.get(st.session_state.generated_image_url)
            image = BytesIO(response.content)

            doc = Document()

            # Option to download content as a Word document
            doc.add_heading(topic, 0)
            doc.add_paragraph(first_line)
            doc.add_picture(image, width=docx.shared.Inches(6))  # Add image to the document
            doc.add_paragraph(remaining_content)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Download as Word Document",
                data=buffer,
                file_name=f"{topic}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
